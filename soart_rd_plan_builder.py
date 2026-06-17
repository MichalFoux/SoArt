from __future__ import annotations

import json
from pathlib import Path

import pdfplumber
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTPUT_DOCX = ROOT / "outputs" / "SoArt_M_RnD_Plan_15062026.docx"


ACCENT = RGBColor(46, 116, 181)
INK = RGBColor(35, 35, 35)
MUTED = RGBColor(95, 95, 95)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_bidi(table) -> None:
    tbl_pr = table._tbl.tblPr
    bidi = tbl_pr.find(qn("w:bidiVisual"))
    if bidi is None:
        bidi = OxmlElement("w:bidiVisual")
        tbl_pr.append(bidi)
    bidi.set(qn("w:val"), "1")


def set_paragraph_rtl(paragraph, align=WD_ALIGN_PARAGRAPH.RIGHT) -> None:
    paragraph.alignment = align
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "1")


def set_run_font(run, name: str = "Arial", size: int | None = None, bold: bool | None = None, color=None) -> None:
    run.font.name = name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for key in ("w:ascii", "w:hAnsi", "w:cs"):
        r_fonts.set(qn(key), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_para(doc: Document, text: str = "", style: str | None = None, bold: bool = False, color=None):
    p = doc.add_paragraph(style=style)
    set_paragraph_rtl(p)
    r = p.add_run(text)
    set_run_font(r, bold=bold, color=color)
    return p


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph(style=f"Heading {level}")
    set_paragraph_rtl(p)
    r = p.add_run(text)
    set_run_font(r, size=16 if level == 1 else 13 if level == 2 else 12, bold=True, color=ACCENT if level < 3 else INK)
    return p


def add_kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    set_table_bidi(table)
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        for idx, cell in enumerate(cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_width(cell, 2700 if idx == 0 else 6660)
            for p in cell.paragraphs:
                set_paragraph_rtl(p)
                for run in p.runs:
                    set_run_font(run, size=10, bold=(idx == 0), color=INK)
            if idx == 0:
                set_cell_shading(cell, "F2F4F7")
    doc.add_paragraph()


def add_matrix_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_bidi(table)
    widths = widths or [9360 // len(headers)] * len(headers)
    for i, head in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = head
        set_cell_shading(cell, "E8EEF5")
        set_cell_width(cell, widths[i])
        for p in cell.paragraphs:
            set_paragraph_rtl(p, WD_ALIGN_PARAGRAPH.CENTER)
            for run in p.runs:
                set_run_font(run, size=9, bold=True, color=INK)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_width(cells[i], widths[i])
            for p in cells[i].paragraphs:
                set_paragraph_rtl(p)
                for run in p.runs:
                    set_run_font(run, size=9, color=INK)
    doc.add_paragraph()


def add_footer(doc: Document) -> None:
    footer = doc.sections[0].footer.paragraphs[0]
    set_paragraph_rtl(footer, WD_ALIGN_PARAGRAPH.CENTER)
    r = footer.add_run("SoArt M R&D Plan | טיוטה לעבודה | 15.06.2026")
    set_run_font(r, size=8, color=MUTED)


def build_soart_docx() -> Path:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style._element.rPr.rFonts.set(qn("w:cs"), "Arial")
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.15
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.color.rgb = ACCENT
    styles["Heading 1"].paragraph_format.space_before = Pt(14)
    styles["Heading 1"].paragraph_format.space_after = Pt(6)
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.color.rgb = ACCENT
    styles["Heading 2"].paragraph_format.space_before = Pt(10)
    styles["Heading 2"].paragraph_format.space_after = Pt(4)

    title = doc.add_paragraph()
    set_paragraph_rtl(title, WD_ALIGN_PARAGRAPH.CENTER)
    run = title.add_run("תכנית מו״פ והטמעה למוצר SoArt M")
    set_run_font(run, size=22, bold=True, color=ACCENT)
    subtitle = doc.add_paragraph()
    set_paragraph_rtl(subtitle, WD_ALIGN_PARAGRAPH.CENTER)
    run = subtitle.add_run("פלטפורמה גלובלית רישתית-דיגיטלית לאמנות ושינוי חברתי")
    set_run_font(run, size=12, color=MUTED)
    meta = doc.add_paragraph()
    set_paragraph_rtl(meta, WD_ALIGN_PARAGRAPH.CENTER)
    run = meta.add_run("טיוטה ראשונית | מותאם מתוך מסמך GLOWE R&D Plan | יוני 2026")
    set_run_font(run, size=9, color=MUTED)

    add_heading(doc, "1. מטרת המסמך", 1)
    add_para(
        doc,
        "מסמך זה ממיר את מבנה תכנית ה-R&D של GLOWE לשפת המוצר, הצרכים והחזון של SoArt M. "
        "במקום תהליך הממוקד בהרשמה של עמותות ומתנדבים, המסמך מגדיר תכנית פיתוח לפלטפורמה "
        "גלובלית המחברת בין אמנים, מחנכים, מטפלים, חוקרים, פעילים חברתיים, מוסדות וקהילות "
        "הפועלים בצומת שבין אמנות, קהילה ושינוי חברתי.",
    )
    add_para(
        doc,
        "המטרה המוצרית היא ליצור תשתית שמאפשרת מעבר מיוזמות מקומיות ומבודדות לאקוסיסטם "
        "מחובר, לומד ומתפתח: רישות, שיתוף ידע, פיתוח יוזמות משותפות, מעקב אימפקט, הכרה מקצועית "
        "והרחבת השפעה ציבורית.",
    )

    add_heading(doc, "2. התאמת הנחות המוצר", 1)
    add_kv_table(
        doc,
        [
            ("מסמך מקור", "GLOWE R&D Plan: שימושי NGO/Volunteer, wishes, opportunities, posts, forums, research וארכיטקטורת שרתים אזורית."),
            ("התאמה ל-SoArt", "SoArt M היא תנועה ופלטפורמה גלובלית לאמנות ושינוי חברתי. המוצר צריך לשרת רשת של אנשים, יוזמות, ידע, קהילות, מוסדות ותהליכי למידה/הטמעה."),
            ("שפת מוצר", "Global Platform, Communities, Education, Healing, Policy & Recognition, Knowledge & Impact."),
            ("עקרון מוביל", "לא רק ניהול פעילות, אלא יצירת שדה פעולה: חיבור בין יצירה, למידה, ידע, פעולה ואימפקט."),
        ],
    )

    add_heading(doc, "3. גרסאות פיתוח מוצעות", 1)
    add_matrix_table(
        doc,
        ["גרסה", "תכלית", "יכולות מרכזיות", "תוצר SoArt"],
        [
            [
                "Initial Release",
                "הקמת שלד מוצרי והגדרת שפה משותפת.",
                "הרשמה, פרופילים בסיסיים, קטגוריות משתמש, עמודי יוזמה, טקסונומיית תחומים, ניהול תוכן ראשוני.",
                "מרחב כניסה שמסביר מי שייך לרשת ומה ניתן לעשות בה.",
            ],
            [
                "V1.0",
                "רישות, גילוי וחיבור בין שחקנים.",
                "חיפוש אנשים/יוזמות/קהילות, מפה או אינדקס תחומי פעילות, קריאות לשיתוף פעולה, פניות חיבור, עדכונים וקהילות עניין.",
                "מעבר מיוזמות מבודדות לאקוסיסטם מחובר.",
            ],
            [
                "V2.0",
                "פיתוח פעולה משותפת, למידה ותיעוד.",
                "מרחבי עבודה ליוזמות, תבניות תוכנית, העלאת ידע, ספריית פרקטיקות, תיעוד תהליכים, מדדי אימפקט בסיסיים, מודרציה.",
                "תנועה לומדת שמייצרת ידע יישומי ולא רק מציגה פעילות.",
            ],
            [
                "V2.1",
                "הרחבה גלובלית והעמקת ידע.",
                "תמיכה רב-לשונית לתוכן ולממשק, מאגר מחקר, חיפוש לפי תגיות ומילות מפתח, שכבת מדיניות והכרה, דוחות אימפקט.",
                "תשתית בינלאומית שמאפשרת חיבור בין תרבויות, מחקר, פעולה והכרה מוסדית.",
            ],
        ],
        [1400, 2200, 3300, 2460],
    )

    add_heading(doc, "4. Use Cases מותאמים ל-SoArt M", 1)
    add_matrix_table(
        doc,
        ["משתמש", "תרחיש שימוש", "ערך מוצרי"],
        [
            ["אמן/ית חברתי/ת", "נרשם/ת, יוצר/ת פרופיל, מציג/ה פרקטיקה או יוזמה ומחפש/ת שותפים.", "נראות, חיבורים מקצועיים והזדמנויות פעולה."],
            ["קהילה או ארגון שטח", "מפרסם צורך, יוזמה, קול קורא או אתגר קהילתי הדורש מענה יצירתי.", "חיבור בין צורך מקומי לבין משאבי יצירה, ידע וליווי."],
            ["מחנך/ת או מטפל/ת", "מאתר/ת מודלים, מערכי פעילות והכשרות בתחום אמנות, חוסן וריפוי.", "למידה יישומית והעברת ידע בין זירות."],
            ["חוקר/ת", "מעלה מחקר, מתייג/ת אותו לפי תחום, ומאתר/ת פרקטיקות או שותפים למחקר פעולה.", "גישור בין ידע אקדמי לבין ידע הנוצר בשדה."],
            ["מנהל/ת SoArt", "מנהל/ת משתמשים, תוכן, מדדים, תגיות, דיווחים ומודרציה.", "שמירה על איכות, אתיקה, שפה משותפת ואמון ברשת."],
            ["שותף מוסדי/פילנתרופי", "סורק יוזמות, תחומי צורך ומדדי אימפקט כדי לזהות הזדמנויות תמיכה.", "קבלת החלטות מבוססת ידע והרחבת השפעה."],
        ],
        [1800, 4400, 3160],
    )

    add_heading(doc, "5. ציר זמן מוצע: 12 שבועות מ-kickoff", 1)
    add_matrix_table(
        doc,
        ["שבוע", "מוקד עבודה", "תוצר"],
        [
            ["1", "יישור חזון, קהלי יעד, מונחים ושפה מוצרית.", "Product brief ו-backlog מאושר."],
            ["2", "ארכיטקטורת מידע: משתמשים, יוזמות, ידע, קהילות ואימפקט.", "מפת ישויות ומסכים ראשיים."],
            ["3", "UX ראשוני ותהליכי onboarding.", "Prototype למסע משתמשים מרכזיים."],
            ["4", "הרשמה, פרופילים והרשאות.", "שלד משתמשים פעיל."],
            ["5", "יוזמות, קהילות וקריאות לשיתוף פעולה.", "מודול פעילות ראשון."],
            ["6", "ספריית ידע, תגיות וחיפוש.", "Knowledge & Practice Library."],
            ["7", "רישות וחיבור בין משתמשים ויוזמות.", "Connection flow והתראות בסיסיות."],
            ["8", "ניהול, מודרציה ומדדי אימפקט ראשוניים.", "Admin + Impact dashboard ראשוני."],
            ["9", "תמיכה רב-לשונית לתוכן ולממשק.", "תשתית תרגום ותצוגה דו/רב-לשונית."],
            ["10", "QA, אבטחת מידע ותיקוני חוויית משתמש.", "גרסת pilot יציבה."],
            ["11", "הזנת תוכן ראשוני מתוך חומרי SoArt וטופז.", "מאגר פתיחה איכותי."],
            ["12", "השקה מוגבלת לקבוצת חלוץ ולמידה.", "Pilot launch + מדדי שימוש ראשונים."],
        ],
        [900, 4300, 4160],
    )

    add_heading(doc, "6. ארכיטקטורה מושגית", 1)
    add_para(
        doc,
        "הארכיטקטורה של SoArt M צריכה לתמוך ברשת גלובלית עם מוקדי פעילות מקומיים. במקום לחשוב רק על "
        "שרתים אזוריים, מומלץ לחשוב על שכבות: שכבת זהות והרשאות; שכבת רשת וקהילות; שכבת יוזמות ופעולה; "
        "שכבת ידע ומחקר; שכבת אימפקט; ושכבת תרגום/לוקליזציה.",
    )
    add_matrix_table(
        doc,
        ["שכבה", "אובייקטים מרכזיים", "הערות פיתוח"],
        [
            ["Identity & Roles", "משתמש, ארגון, תפקיד, תחום מומחיות, אזור פעולה.", "תמיכה בתפקידים חופפים: אמן יכול להיות גם מחנך, חוקר או מוביל קהילה."],
            ["Network & Communities", "קהילה, תחום עניין, מרחב פעולה, קשרים, הזמנות.", "בניית רשת של רשתות ולא רשימת משתמשים שטוחה."],
            ["Initiatives & Opportunities", "יוזמה, צורך, קול קורא, תוכנית, סטטוס, שותפים.", "אפשרות לעבור מרעיון לשיתוף פעולה ולהטמעה בשטח."],
            ["Knowledge & Research", "מאמר, מודל, כלי, מקרה בוחן, תיעוד, תגיות.", "חיפוש חופשי, תגיות ושיוך לתחומי פעילות של SoArt."],
            ["Impact & Learning", "מדד, עדות, תוצאה, אוכלוסיית יעד, למידה.", "מדידה רכה וכמותית: חוסן, שייכות, השתתפות, שינוי תפיסתי והשפעה ציבורית."],
            ["Localization", "שפה, תרגום, הקשר תרבותי, אזור.", "נדרש כבר מ-V2.1 כדי לאפשר תנועה גלובלית אמיתית."],
        ],
        [1900, 3300, 4160],
    )

    add_heading(doc, "7. מדדי הצלחה ראשוניים", 1)
    add_matrix_table(
        doc,
        ["תחום", "מדדים מוצעים"],
        [
            ["רשת", "מספר משתמשים פעילים, קשרים חדשים, קהילות פעילות, שיעור השלמת פרופיל."],
            ["ידע", "פריטי ידע שהועלו, חיפושים, הורדות/שמירות, שימוש חוזר בפרקטיקות."],
            ["פעולה", "יוזמות שנפתחו, שותפויות שנוצרו, קריאות שנענו, תוכניות שהגיעו להטמעה."],
            ["אימפקט", "עדויות שדה, מדדי השתתפות, חוסן ושייכות, דוחות למידה, השפעה על מדיניות/הכרה."],
            ["איכות ואמון", "תוכן שעבר אוצרות, אירועי מודרציה, שביעות רצון משתמשים, שמירה על אתיקה ופרטיות."],
        ],
        [2200, 7160],
    )

    add_heading(doc, "8. החלטות פתוחות להמשך", 1)
    add_matrix_table(
        doc,
        ["נושא", "החלטה נדרשת"],
        [
            ["שם ומיתוג מוצר", "האם להשתמש ב-SoArt M, SoArt Movement או שם מוצר נפרד לפלטפורמה הדיגיטלית."],
            ["קהלי יעד בשלב pilot", "האם להתחיל באמנים ויוזמות שטח, במוסדות, בחוקרים או בקבוצת חלוץ מעורבת."],
            ["שפות השקה", "עברית ואנגלית בלבד בשלב ראשון, או הוספת שפות יעד נוספות כבר ב-pilot."],
            ["רמת פתיחות", "מה פתוח לציבור, מה דורש הרשמה, ומה מוגבל לקהילות/שותפים מאושרים."],
            ["מדידת אימפקט", "אילו מדדים בסיסיים נאספים בלי להכביד על קהילות ויוזמות בשטח."],
        ],
        [2400, 6960],
    )

    add_heading(doc, "9. המלצת מיקוד לשלב הראשון", 1)
    add_para(
        doc,
        "מומלץ שהגרסה הראשונה לא תנסה לכסות את כל חזון התנועה, אלא תבנה שלד עובד שמוכיח את "
        "הערך המרכזי: חיבור בין אנשים, יוזמות וידע בשדה האמנות לשינוי חברתי. לכן ה-MVP צריך לכלול "
        "פרופילים, יוזמות, ספריית ידע בסיסית, חיפוש, קריאות לשיתוף פעולה, וממשק ניהול איכותי. "
        "רק לאחר שיש תנועה אמיתית בתוך המערכת כדאי להעמיק למדדי אימפקט, רב-לשוניות מלאה, מחקר מתקדם "
        "ודוחות מדיניות.",
    )

    add_footer(doc)
    OUTPUT_DOCX.parent.mkdir(exist_ok=True)
    doc.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


def extract_pdf_text(pdf_path: Path) -> list[str]:
    pages: list[str] = []
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            pages.append(page.extract_text() or "")
    return pages


def extract_docx_text(docx_path: Path, limit: int | None = None) -> list[str]:
    doc = Document(docx_path)
    lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return lines if limit is None else lines[:limit]


def main() -> None:
    data = {
        "glowe_pdf_pages": extract_pdf_text(ROOT / "CF31120626 GLOWE R&D plan.pdf"),
        "soart_sources": {
            "abstract": extract_docx_text(ROOT / "אבסטרקט מורחב SoArt.docx", 120),
            "short_prospectus": extract_docx_text(ROOT / "תנועת סוארט תשקיף קצר 03052026.docx", 160),
            "landing_page": extract_docx_text(ROOT / "לאתר טופז בינלאומי דף נחיתה סוארט.docx", 100),
        },
    }
    out = ROOT / "outputs" / "soart_extracted_context.json"
    out.parent.mkdir(exist_ok=True)
    out.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    print("outputs/soart_extracted_context.json")
    print("glowe_pages", len(data["glowe_pdf_pages"]))
    for i, page in enumerate(data["glowe_pdf_pages"], 1):
        print(f"page_{i}_chars", len(page))
    for name, lines in data["soart_sources"].items():
        print(name, len(lines))
    docx_path = build_soart_docx()
    print("created", docx_path.name)


if __name__ == "__main__":
    main()
